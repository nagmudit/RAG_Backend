import os
import logging
from typing import List, Dict, Any
from pathlib import Path
import aiofiles
from io import BytesIO

# Document processing libraries
import PyPDF2
from docx import Document
import openpyxl
import markdown
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of various document formats for RAG."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        
    async def process_uploaded_file(self, file_content: bytes, filename: str) -> List[LangchainDocument]:
        """
        Process an uploaded file and return LangChain documents.
        
        Args:
            file_content: The raw bytes of the uploaded file
            filename: The name of the uploaded file
            
        Returns:
            List of LangChain Document objects
        """
        try:
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.pdf':
                text = await self._extract_pdf_text(file_content)
            elif file_extension == '.docx':
                text = await self._extract_docx_text(file_content)
            elif file_extension == '.xlsx':
                text = await self._extract_xlsx_text(file_content)
            elif file_extension == '.md':
                text = await self._extract_markdown_text(file_content)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            if not text.strip():
                raise ValueError(f"No text content found in {filename}")
            
            # Create metadata
            metadata = {
                "source": filename,
                "file_type": file_extension,
                "title": Path(filename).stem
            }
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create LangChain documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = metadata.copy()
                doc_metadata["chunk"] = i
                doc_metadata["total_chunks"] = len(chunks)
                
                documents.append(LangchainDocument(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            logger.info(f"Processed {filename}: {len(documents)} chunks created")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise ValueError(f"Failed to process {filename}: {str(e)}")
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_xlsx_text(self, file_content: bytes) -> str:
        """Extract text from XLSX file."""
        try:
            xlsx_file = BytesIO(file_content)
            workbook = openpyxl.load_workbook(xlsx_file, read_only=True)
            
            text = ""
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                
                for row in worksheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None and str(cell_value).strip():
                            row_text.append(str(cell_value).strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            workbook.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting XLSX text: {e}")
            raise ValueError(f"Failed to extract text from XLSX: {str(e)}")
    
    async def _extract_markdown_text(self, file_content: bytes) -> str:
        """Extract text from Markdown file."""
        try:
            # Decode bytes to string
            markdown_content = file_content.decode('utf-8')
            
            # Convert Markdown to HTML, then extract text
            html = markdown.markdown(markdown_content)
            
            # For now, we'll keep the original markdown content
            # as it's already in a readable text format
            return markdown_content.strip()
            
        except UnicodeDecodeError:
            try:
                # Try different encoding
                markdown_content = file_content.decode('latin-1')
                return markdown_content.strip()
            except Exception as e:
                logger.error(f"Error decoding markdown file: {e}")
                raise ValueError(f"Failed to decode markdown file: {str(e)}")
        except Exception as e:
            logger.error(f"Error extracting Markdown text: {e}")
            raise ValueError(f"Failed to extract text from Markdown: {str(e)}")

# Global instance
document_processor = DocumentProcessor()
